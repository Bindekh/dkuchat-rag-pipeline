
import os
import pickle
from pathlib import Path
from collections import Counter
 
import pdfplumber
import docx
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
from rank_bm25 import BM25Okapi
 
# Optional OCR 
# if the PDF is text-based.
try:
    import pytesseract
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
 
# ── Config ──────────────────────────────────────────────────────────────────
DOCS_DIR      = "data/"
INDEX_DIR     = "index/"
CHUNK_SIZE    = 400   # words per chunk
CHUNK_OVERLAP = 80    # word overlap between chunks
MIN_WORDS     = 30    # discard tiny tail fragments
 
 
# ── Core chunker ────────────────────────────────────────────────────────────
 
def _slide(words, source, page, chunks):
    """Append sliding-window chunks to the shared list."""
    step = CHUNK_SIZE - CHUNK_OVERLAP
    for i in range(0, len(words), step):
        cw = words[i : i + CHUNK_SIZE]
        if len(cw) < MIN_WORDS:
            continue
        chunks.append({
            "text":     " ".join(cw),
            "source":   source,
            "page":     page,
            "chunk_id": len(chunks),
        })
 
 
# ── PDF loader ───────────────────────────────────────────────────────────────
 
def _ocr_pdf(file_path):
    """
    Convert every page to an image and run Tesseract OCR.
    Returns a list of (page_num, text) tuples.
    """
    if not OCR_AVAILABLE:
        print("  [OCR] pytesseract / pdf2image not installed — skipping OCR.")
        print("        Run:  brew install tesseract && pip install pytesseract pdf2image")
        return []
 
    print(f"  [OCR] Scanned PDF detected — running Tesseract on {file_path.name} …")
    images = convert_from_path(str(file_path), dpi=200)
    results = []
    for page_num, img in enumerate(images, start=1):
        text = pytesseract.image_to_string(img, lang="eng")
        if text.strip():
            results.append((page_num, text))
        if page_num % 50 == 0:
            print(f"  [OCR]   … page {page_num}/{len(images)}")
    return results
 
 
def _is_scanned_pdf(file_path, probe_pages=5):
    """Return True if first probe_pages pages all have no extractable text."""
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            pages_to_check = min(probe_pages, len(pdf.pages))
            for i in range(pages_to_check):
                page = pdf.pages[i]
                text = (page.extract_text() or "").strip()
                if not text:
                    text = " ".join(w["text"] for w in page.extract_words()).strip()
                if text:
                    return False
        return True
    except Exception:
        return True
 
 
def _load_pdf(file_path, chunks):
    """
    Extract text page-by-page.
    - Scanned PDF (image-based): skip pdfplumber, go straight to OCR.
    - Text-based PDF: pdfplumber with extract_words() fallback + table extraction.
    """
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            npages = len(pdf.pages)
        print(f"  [PDF] {file_path.name}: {npages} pages")
 
        # Scanned PDF fast-path
        if _is_scanned_pdf(file_path):
            print(f"  [PDF] No extractable text found — using OCR")
            ocr_results = _ocr_pdf(file_path)
            extracted = 0
            for page_num, text in ocr_results:
                before = len(chunks)
                _slide(text.split(), file_path.name, page_num, chunks)
                extracted += len(chunks) - before
            print(f"  [PDF] -> {extracted} chunks")
            return
 
        # Text-based PDF normal path
        extracted   = 0
        image_pages = 0
        with pdfplumber.open(str(file_path)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if not text.strip():
                    text = " ".join(w["text"] for w in page.extract_words())
                for table in (page.extract_tables() or []):
                    for row in table:
                        row_text = " | ".join(
                            str(cell).strip() for cell in row if cell
                        )
                        if row_text:
                            text += "\n" + row_text
                words = text.split()
                if not words:
                    image_pages += 1
                    continue
                before = len(chunks)
                _slide(words, file_path.name, page_num, chunks)
                extracted += len(chunks) - before
 
        if image_pages:
            print(f"  [PDF] {image_pages} image-only pages skipped")
        print(f"  [PDF] -> {extracted} chunks")
 
    except Exception as e:
        print(f"  [PDF] ERROR reading {file_path.name}: {e}")
 
 
# ── DOCX loader ──────────────────────────────────────────────────────────────
 
def _load_docx(file_path, chunks):
    """
    Extract text from a DOCX file with per-page tracking.
 
    Strategy:
      1. Collect all paragraph and table text in order.
      2. Detect page breaks via XML <w:pageBreak/> elements.
      3. If no breaks found, split into logical pages by word-count windows
         (~500 words per page).
    """
    try:
        doc = docx.Document(str(file_path))
 
        PAGEBREAK_TAG = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pageBreak"
        pages_text = []
        current    = []
 
        def flush_page():
            text = "\n".join(current).strip()
            if text:
                pages_text.append(text)
            current.clear()
 
        for para in doc.paragraphs:
            if para._element.find(f".//{PAGEBREAK_TAG}") is not None:
                flush_page()
            txt = para.text.strip()
            if txt:
                current.append(txt)
 
        # Table content appended after paragraphs
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    current.append(row_text)
 
        flush_page()
 
        # If no page breaks detected, split by word-count window
        if len(pages_text) == 1:
            all_words    = pages_text[0].split()
            words_per_pg = 500
            pages_text   = []
            for start in range(0, len(all_words), words_per_pg):
                seg = " ".join(all_words[start : start + words_per_pg])
                if seg.strip():
                    pages_text.append(seg)
 
        print(f"  [DOCX] {file_path.name}: {len(pages_text)} logical pages")
        extracted = 0
        for page_num, text in enumerate(pages_text, start=1):
            before = len(chunks)
            _slide(text.split(), file_path.name, page_num, chunks)
            extracted += len(chunks) - before
 
        print(f"  [DOCX] -> {extracted} chunks")
    except Exception as e:
        print(f"  [DOCX] ERROR reading {file_path.name}: {e}")
 
 
# ── Main loader ──
 
def load_and_chunk_files(docs_dir):
    chunks = []
    files  = sorted(Path(docs_dir).glob("*"))
 
    if not files:
        print(f"No files found in {docs_dir}")
        return chunks
 
    for file_path in files:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            _load_pdf(file_path, chunks)
        elif suffix == ".docx":
            _load_docx(file_path, chunks)
        else:
            print(f"  [SKIP] {file_path.name} (unsupported type)")
 
    # Re-number chunk_ids sequentially
    for i, c in enumerate(chunks):
        c["chunk_id"] = i
 
    return chunks
 
 
# ── Index builders ───────────────────────────────────────────────────────────
 
def build_vector_index(chunks, model_name, index_dir):
    model      = SentenceTransformer(model_name)
    texts      = [c["text"] for c in chunks]
    embeddings = model.encode(texts, show_progress_bar=True, batch_size=32)
    embeddings = np.array(embeddings).astype("float32")
    faiss.normalize_L2(embeddings)
 
    index = faiss.IndexFlatIP(embeddings.shape[1])
    index.add(embeddings)
 
    os.makedirs(index_dir, exist_ok=True)
    safe = model_name.replace("/", "_")
    faiss.write_index(index, f"{index_dir}/faiss_{safe}.index")
    with open(f"{index_dir}/chunks.pkl", "wb") as f:
        pickle.dump(chunks, f)
 
    print(f"Built vector index: {len(chunks)} chunks  [{model_name}]")
 
 
def build_bm25_index(chunks, index_dir):
    tokenized = [c["text"].lower().split() for c in chunks]
    bm25      = BM25Okapi(tokenized)
    with open(f"{index_dir}/bm25.pkl", "wb") as f:
        pickle.dump(bm25, f)
    print("Built BM25 index")
 
 
# ── Entry point ──────────────────────────────────────────────────────────────
 
if __name__ == "__main__":
    EMBEDDING_MODEL = "BAAI/bge-small-en-v1.5"
 
    print(f"Loading files from: {DOCS_DIR}")
    chunks = load_and_chunk_files(DOCS_DIR)
    print(f"\nTotal chunks: {len(chunks)}")
 
    if not chunks:
        print("No chunks found! Make sure your files are in the data/ folder.")
        print("  If PDF pages are image-based, add pytesseract OCR.")
    else:
        counts = Counter(c["source"] for c in chunks)
        for src, n in sorted(counts.items()):
            print(f"  {src}: {n} chunks")
 
        build_vector_index(chunks, EMBEDDING_MODEL, INDEX_DIR)
        build_bm25_index(chunks, INDEX_DIR)
 